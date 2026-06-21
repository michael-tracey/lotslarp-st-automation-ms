#!/usr/bin/env python3
"""
LOTS LARP — Player Nominations Doc Builder
Reads the "Nominations?" column from a monthly sheet, cleans up the raw text
(drops blanks, strips leading hyphens/bullets and stray quotes), randomizes the
order, and writes a local .docx file titled:

    LOTSLARP Player Nominations: <Month> <Year>

…with every nomination as a bullet point. Upload the .docx to Google Drive and
open it (or File → Open) to convert it to a native Google Doc — the heading and
bullets carry over.

This replaces the old command-line dance:
    awk '!/^$/' june.txt | shuf | sed "s/^\\-//"

Reading the sheet uses the same service account as the other utilities.

Setup:
  pip install -r requirements.txt
  cp .env-example .env             # fill in your values

Usage:
  python create_nominations_doc.py                 # pick the month interactively
  python create_nominations_doc.py "June 2026"     # skip the month picker
  python create_nominations_doc.py -o ~/noms.docx  # choose the output path
  python create_nominations_doc.py --dry-run       # print the cleaned list, write nothing
  python create_nominations_doc.py --no-shuffle    # keep sheet order instead of randomizing
"""

import argparse
import os
import random
import re
import sys
from datetime import datetime

from rich.console import Console
from rich.panel import Panel

from gspread.utils import rowcol_to_a1

sys.path.insert(0, os.path.dirname(__file__))
from config import (
    SPREADSHEET_ID, MONTH_YEAR_PATTERN,
    sheets_client, sheets_retry,
)

console = Console()


# ── Text cleaning ───────────────────────────────────────────────────────────────

_LEADING_MARKERS = re.compile(r'^[\s\-–—•·*>]+')
_DOUBLE_QUOTES   = ('"', '“', '”')   # " “ ”
_WRAP_PAIRS      = (('"', '"'), ("'", "'"), ('“', '”'), ('‘', '’'))


def _strip_wrapping_quotes(s):
    """Remove a matched pair of quotes wrapping the whole string (straight or
    smart), repeatedly, e.g.  "hello"  →  hello."""
    changed = True
    while changed and len(s) >= 2:
        changed = False
        for open_q, close_q in _WRAP_PAIRS:
            if s[0] == open_q and s[-1] == close_q:
                inner = s[1:-1].strip()
                if inner:
                    s, changed = inner, True
                    break
    return s


def _strip_unmatched_double(s):
    """If there's an odd number of double quotes, one is dangling — drop the
    leading one (preferred) or the trailing one. Handles the common
    `"I couldn't attend` case without mangling apostrophes/contractions."""
    if sum(s.count(q) for q in _DOUBLE_QUOTES) % 2 == 1:
        if s[:1] in _DOUBLE_QUOTES:
            s = s[1:].strip()
        elif s[-1:] in _DOUBLE_QUOTES:
            s = s[:-1].strip()
    return s


def clean_nomination(raw):
    """Clean a single nomination line. Returns '' if nothing meaningful remains."""
    s = raw.replace('\r', '').strip()
    if not s:
        return ''
    s = _LEADING_MARKERS.sub('', s).strip()   # leading hyphens / bullets
    s = _strip_wrapping_quotes(s)             # "fully wrapped" → unwrapped
    s = _strip_unmatched_double(s)            # dangling leading/trailing quote
    s = re.sub(r'[ \t]{2,}', ' ', s).strip()  # collapse runs of spaces
    return s


def extract_nominations(column_cells):
    """Each cell may hold several nominations on separate lines. Split them out,
    clean each, drop blanks. Returns a flat list."""
    out = []
    for cell in column_cells:
        for line in cell.split('\n'):
            cleaned = clean_nomination(line)
            if cleaned:
                out.append(cleaned)
    return out


# ── Sheet / column pickers ──────────────────────────────────────────────────────

def list_month_sheets(spreadsheet):
    candidates = [ws for ws in spreadsheet.worksheets()
                  if MONTH_YEAR_PATTERN.match(ws.title)]

    def sort_key(ws):
        try:
            return datetime.strptime(ws.title, "%B %Y")
        except ValueError:
            return datetime.min

    return sorted(candidates, key=sort_key, reverse=True)


def pick_sheet(spreadsheet):
    sheets = list_month_sheets(spreadsheet)
    if not sheets:
        console.print("[red]No 'Month Year' sheets found in the spreadsheet.[/red]")
        sys.exit(1)

    console.print("\n[bold]Available monthly sheets:[/bold]")
    for i, ws in enumerate(sheets):
        tag = "  [dim]← most recent (default)[/dim]" if i == 0 else ""
        console.print(f"  [cyan]{i + 1}.[/cyan] {ws.title}{tag}")

    while True:
        try:
            raw = input("\nSelect sheet [1]: ").strip()
        except (EOFError, KeyboardInterrupt):
            console.print()
            sys.exit(0)
        if not raw:
            return sheets[0].title
        try:
            idx = int(raw) - 1
            if 0 <= idx < len(sheets):
                return sheets[idx].title
        except ValueError:
            pass
        console.print(f"  [yellow]Enter a number between 1 and {len(sheets)}.[/yellow]")


def _col_letter(idx_zero_based):
    return rowcol_to_a1(1, idx_zero_based + 1).rstrip('1')


def detect_nominations_column(headers):
    """Return the 0-based index of the Nominations column, or None."""
    for i, h in enumerate(headers):
        if h.strip().lower().rstrip('?') == 'nominations':
            return i
    for i, h in enumerate(headers):           # looser fallback
        if 'nomination' in h.strip().lower():
            return i
    return None


def pick_column(headers):
    """Auto-detect the 'Nominations?' column and let the user confirm or override.
    Returns the 0-based column index."""
    detected = detect_nominations_column(headers)

    if detected is not None:
        col = _col_letter(detected)
        console.print(
            f"\nDetected column [bold]{col}[/bold] "
            f"— [cyan]\"{headers[detected]}\"[/cyan]"
        )
        try:
            raw = input("Use this column? [Y/n, or a column letter]: ").strip()
        except (EOFError, KeyboardInterrupt):
            console.print()
            sys.exit(0)
        if not raw or raw.lower() in ('y', 'yes'):
            return detected
        # A bare column letter is a quick override
        override = _resolve_column_letter(raw, headers)
        if override is not None:
            return override
    else:
        console.print("\n[yellow]Couldn't find a 'Nominations?' header.[/yellow]")

    # Fall back to listing every header for an explicit pick
    return _choose_from_headers(headers)


def _resolve_column_letter(raw, headers):
    raw = raw.strip().upper()
    if not re.fullmatch(r'[A-Z]+', raw):
        return None
    idx = 0
    for ch in raw:
        idx = idx * 26 + (ord(ch) - ord('A') + 1)
    idx -= 1
    if 0 <= idx < len(headers):
        return idx
    console.print(f"  [yellow]Column {raw} is out of range.[/yellow]")
    return None


def _choose_from_headers(headers):
    console.print("\n[bold]Columns:[/bold]")
    for i, h in enumerate(headers):
        label = h.strip() or "[dim](blank)[/dim]"
        console.print(f"  [cyan]{_col_letter(i)}[/cyan]  {label}")
    while True:
        try:
            raw = input("\nEnter the column letter to read: ").strip()
        except (EOFError, KeyboardInterrupt):
            console.print()
            sys.exit(0)
        idx = _resolve_column_letter(raw, headers)
        if idx is not None:
            return idx
        console.print("  [yellow]Enter a valid column letter (e.g. AH).[/yellow]")


# ── .docx output ────────────────────────────────────────────────────────────────

def _safe_filename(name):
    """Turn a doc title into a filesystem-friendly base name."""
    return re.sub(r'[\\/:*?"<>|]+', '', name).strip()


def write_docx(title, nominations, path):
    """Write a .docx with a Heading 1 title and a bulleted list. Uploading this
    file to Google Drive and opening it produces a native Google Doc with the
    heading and bullets preserved."""
    try:
        from docx import Document
    except ImportError:
        console.print(
            "[red]Missing python-docx.[/red] Run: "
            "[bold]pip install -r requirements.txt[/bold]"
        )
        sys.exit(1)

    document = Document()
    document.add_heading(title, level=1)
    for nomination in nominations:
        document.add_paragraph(nomination, style='List Bullet')
    document.save(path)


# ── Main ────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="LOTS LARP Player Nominations Doc Builder")
    parser.add_argument('sheet', nargs='?',
                        help='Sheet name e.g. "June 2026" — skips the interactive picker')
    parser.add_argument('-o', '--output',
                        help='Output .docx path (default: "LOTSLARP Player Nominations - '
                             '<Month> <Year>.docx" in the current directory)')
    parser.add_argument('--dry-run', action='store_true',
                        help='Print the cleaned, randomized list to the console; write no file')
    parser.add_argument('--no-shuffle', action='store_true',
                        help='Keep the order from the sheet instead of randomizing')
    args = parser.parse_args()

    console.print("Connecting to Google Sheets…")
    client      = sheets_client()
    spreadsheet = client.open_by_key(SPREADSHEET_ID)

    sheet_name = args.sheet if args.sheet else pick_sheet(spreadsheet)
    console.print(f"Selected: [bold]{sheet_name}[/bold]")

    sheet      = spreadsheet.worksheet(sheet_name)
    all_values = sheets_retry(
        sheet.get_all_values,
        notify=lambda msg: console.print(f"  [yellow]{msg}[/yellow]"),
    )
    if not all_values:
        console.print("[red]Sheet is empty.[/red]")
        sys.exit(1)

    headers = all_values[0]
    col_idx = pick_column(headers)
    console.print(
        f"Reading column [bold]{_col_letter(col_idx)}[/bold] "
        f"(\"{headers[col_idx]}\")"
    )

    column_cells = [row[col_idx] if col_idx < len(row) else '' for row in all_values[1:]]
    nominations  = extract_nominations(column_cells)

    if not nominations:
        console.print("[yellow]No nominations found in that column.[/yellow]")
        sys.exit(0)

    if not args.no_shuffle:
        random.shuffle(nominations)

    console.print(f"  [bold]{len(nominations)}[/bold] nominations after cleanup"
                  f"{'' if args.no_shuffle else ', randomized'}.")

    title = f"LOTSLARP Player Nominations: {sheet_name}"

    if args.dry_run:
        body = "\n".join(f"• {n}" for n in nominations)
        console.print(Panel(body, title=f"[bold]{title}[/bold]  [dim](dry run)[/dim]",
                            border_style="dim", expand=False))
        console.print("\n[yellow]Dry run — no file written.[/yellow]")
        return

    out_path = args.output or f"{_safe_filename(title)}.docx"
    write_docx(title, nominations, out_path)
    console.print(f"\n[green]✓ Wrote:[/green] [bold]{os.path.abspath(out_path)}[/bold]")
    console.print("  [dim]Upload it to Google Drive and open it to convert to a "
                  "Google Doc (heading + bullets carry over).[/dim]")


if __name__ == "__main__":
    main()
